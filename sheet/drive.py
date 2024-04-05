import docx
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import os
from .models import FileName

# from download import file_path




# directory = 'D:\\python\\google sheet\\sheet'
#
# for i in os.listdir(directory):
#     filename = os.path.join(directory, i)
#     gfile = drive.CreateFile({"parents": [{"id": folder}], "title": i})
#     gfile.SetContentFile(filename)
#     gfile.Upload()
def download():
    # gauth = GoogleAuth()
    # gauth.LocalWebserverAuth()

    drive = GoogleDrive()

    folder = '1eUj7ys6Mh7EEoPPVCqH9SsMRLscJBvYF'
    file_list = drive.ListFile({"q": f"'{folder}' in parents and trashed=false"}).GetList()
    list_file = []
    for index, file in enumerate(file_list):
        print(index + 1, "file downloaded: ", file['title'])
        file.GetContentFile(file['title'])
        list_file.append(file['title'])

    # print("List", list_file)

    folder_path = os.getcwd()
    # print(folder_path)

    new_data = {
        "C_KEY": "Hello Amerika00",
        "D_KEY": "Hello Paris00",
        # Add more key-word pairs as needed
    }

    def update_docx_files_in_folder(folder_path, new_data):
        for filename in list_file:
            if filename.endswith('.docx'):
                file_path = os.path.join(folder_path, filename)
                file_name = os.path.splitext(filename)[0]
                # update_docx_file(file_path, new_data, file_name)  # Fayl nomini ham uzatish
                writing_sheet(file_path, file_name)

    # writing_file = []
    # writing_key = []
    # writing_word = []

    def writing_sheet(file_path, file_name):
        # print(file_path)
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            word1 = paragraph.text.split(":")

            # print(word1)
            # writing_file.append(file_name)
            if word1[0] == "":
                f = FileName(name=file_name, key=word1[0])
                f.save()
            else:
                f = FileName(name=file_name, key=word1[0], word=word1[1])
                f.save()
        # print(writing_file)
        doc.save(file_path)

    update_docx_files_in_folder(folder_path, new_data)
    # writing(writing_file, writing_key, writing_word)

# def update_docx_file(file_path, new_data, file_name):
#     doc = docx.Document(file_path)
#     for paragraph in doc.paragraphs:
#         word1 = paragraph.text.split(":")
#         for key, word in new_data.items():
#             if key in word1[0]:
#                 # updated_paragraphs.append((file_name, word1[0], word))  # Fayl nomini ham qo'shish
#                 file_list = drive.ListFile({"q": f"'{folder}' in parents and trashed=false"}).GetList()
#                 for file in file_list:
#                     if file['title'] == file_name + ".docx":
#                         file.SetContentString(word1[0] + ":" + word)
#                         file.Upload()
#                         print(file['title'])
#                         paragraph.text = word1[0] + ":" + word
#     doc.save(file_path)
#
#
# update_docx_files_in_folder(folder_path, new_data)
